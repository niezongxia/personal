#!/usr/bin/env python
# -*- coding: utf-8 -*-
import os
import re
import sys
import time
import docx
import json
import requests
import datetime
import shutil
from openpyxl import Workbook
from openpyxl.reader.excel import load_workbook
from docxcompose.composer import Composer
from docx import Document as Document_compose
from docx.enum.text import (WD_ALIGN_PARAGRAPH,WD_TAB_ALIGNMENT,WD_TAB_LEADER,WD_LINE_SPACING,WD_COLOR_INDEX)
from docx.shared import (Length,Inches,Pt, Cm,Mm,RGBColor,)

############## 倒计时函数 #########################
def Windows_close(close_times):
    for x in range(close_times,-1,-1):
        mystr = "程序执行完毕！窗口将在" + str(x) + "秒后关闭！！！"
        print(mystr,end = "")
        print("\b" * (len(mystr)*2),end = "",flush = True)
        time.sleep(1)

################# 用户自定义信息 ##############
excel_name = sys.argv[1]#指定预置表路径
excel_check = sys.argv[2]
title1 = sys.argv[3]
title2 = sys.argv[4]
title3 = sys.argv[5]
title4 = sys.argv[6]
title5 = sys.argv[7]
title6 = sys.argv[8]
tester = sys.argv[9]#'高波'
devlpor = sys.argv[10]#'陈雷'
teamer = sys.argv[11]#'林鹏鹏'
task = sys.argv[12]# 'MLIFE-APPS_V01.00_B339'

############# 自动获取及目录写死 #############
home = os.getcwd().replace('\\','/') + '/'#获取当前工作目录
model = 'model'
file_source = 'file_source'
install_book = 'install_book'

######## 根路径写死参数 ###############
src = '/src'
exe = '/exe'
indoc = '/indoc'
reldoc = '/reldoc'

#################### 后台包路径 ##################
war = exe + '/war'
sql = exe + '/sql'
sql_bak = exe + '/备份'
sql_back = exe + '/回滚'
sql_select = exe + '/验证'

######### 客户端包路径 #########
phone = exe+'/apk'

######## 目录拼接 ################
home_task = home + task
task_src = home_task + src
task_exe = home_task + exe
task_indoc = home_task + indoc
task_reldoc = home_task + reldoc
task_war = home_task + war
task_sql = home_task + sql
task_bak = home_task + sql_bak
task_back = home_task + sql_back
task_select = home_task + sql_select
task_phone = home_task + phone
package_model = home + model
package_install_book = home + install_book
excel_path = package_install_book + '/' + excel_name
excel_check_path = package_install_book + '/' + excel_check
code_dir = task_exe+"/"

######## 定义列表 ##############
APPS_dir_list = [home_task,task_reldoc]
Phone_dir_list = [home_task,task_reldoc]
model_list = []
model_check_list = {}
model_use_list = {}
check_result = []
result=[]
exe_list = []
src_list = []
code_list = ['sql','war','回滚']
slist = []
rlist = []
plist = []
wlist = ['head']
HY_check_list = []
HSH_check_list = []
HY_combine_list = []
HSH_combine_list = []
############## 获取分类匹配信息 ##################
build_path = package_model + '/build.txt'
build_paths={}
with open(build_path,'r',encoding = 'UTF-8') as df:
    for kv in [d.strip().split(' = ') for d in df]:
        build_paths[kv[0]]=kv[1]

files_auto_exe = u''.join(build_paths.get('files_auto_exe'))
files_checklist = u''.join(build_paths.get('files_checklist'))
files_auto_ins = u''.join(build_paths.get('files_auto_ins'))
files_indoc = u''.join(build_paths.get('files_indoc'))
files_reldoc = u''.join(build_paths.get('files_reldoc'))
files_war = u''.join(build_paths.get('files_war'))
files_sql = u''.join(build_paths.get('files_sql'))
return_sql = u''.join(build_paths.get('return_sql'))
bak_sql = u''.join(build_paths.get('bak_sql'))
select_sql = u''.join(build_paths.get('select_sql'))
files_src = u''.join(build_paths.get('files_src'))
file_phone = u''.join(build_paths.get('file_phone'))

############## 提示打包类型 ###############
if re.search('_B',task) !=None:
    print('测试包开始组包...')
elif re.search('_SP',task) !=None:
    print('生产包开始组包...')
elif re.search('_PTF',task) !=None:
    print('解bug生产包开始组包...')

########### 判断组包类型 ##############
if re.search('MLIFE-PHONE',task) !=None:
    dir_list = Phone_dir_list
    task_type = 'MLIFE-PHONE'
elif re.search('MLIFE-APPS',task) !=None:
    dir_list = APPS_dir_list
    task_type = 'MLIFE-APPS'
elif re.search('DSMS-VABS',task) !=None:
    dir_list = APPS_dir_list
    task_type = 'DSMS-VABS'
elif re.search('BCSP-DATR',task) !=None:
    dir_list = APPS_dir_list
    task_type = 'BCSP-DATR'
elif re.search('BCSP-DASP',task) !=None:
    dir_list = APPS_dir_list
    task_type = 'BCSP-DASP'
elif re.search('BCSP-CCST',task) !=None:
    dir_list = APPS_dir_list
    task_type = 'BCSP-CCST'
elif re.search('BCSP-UPSMS',task) !=None:
    dir_list = APPS_dir_list
    task_type = 'BCSP-UPSMS'

################## 开始组建基础目录 ##############
for file_dir in dir_list:
    if os.path.exists(file_dir) is False:
        os.makedirs(file_dir)

################## 构建文件清单list #####################
def find_file(file_list,file_dir,file_mark_list):
    if os.path.exists(file_dir) == True:
        if file_mark_list !=[]:
            for file_mark in file_mark_list:
                wlist.append(file_mark)
        sl_s = list(os.walk(file_dir))
        for si in sl_s:
            os.chdir(si[0])#进入本级路径，防止找不到文件而报错
            if si[2]!=[]:#如该路径下有文件
                for x in si[2]:#遍历文件
                    file_list.append(x)
            else:
                print('没有执行码！')

####################### 构建扫描路径清单 ####################
def fileListFunc(filePathList,exe_list,src_list):
    for filePath in filePathList:
        i = 0
        for top, dirs, nondirs in os.walk(filePath):
            for item in nondirs:
                home_path = os.path.join(top, item).replace('\\','/')
                ### 判断程序
                if re.search(files_indoc,item) != None:
                    if os.path.exists(task_indoc):
                        task_path = os.path.join(task_indoc, item).replace('\\','/')
                    else:
                        os.makedirs(task_indoc)
                        task_path = os.path.join(task_indoc, item).replace('\\','/')
                elif re.search(files_reldoc,item) != None:
                    if os.path.exists(task_reldoc):
                        task_path = os.path.join(task_reldoc, item).replace('\\','/')
                    else:
                        os.makedirs(task_reldoc)
                        task_path = os.path.join(task_reldoc, item).replace('\\','/')
                elif re.search(files_src,item) != None:
                    if os.path.exists(task_src):
                        task_path = os.path.join(task_src, item).replace('\\','/')
                    else:
                        os.makedirs(task_src)
                        task_path = os.path.join(task_src, item).replace('\\','/')
                    src_list.append(os.path.join(src, item).replace('/','\\').replace('\src','src'))
                elif re.search(files_war,item) != None:
                    if os.path.exists(task_war):
                        task_path = os.path.join(task_war, item).replace('\\','/')
                    else:
                        os.makedirs(task_war)
                        task_path = os.path.join(task_war, item).replace('\\','/')
                    exe_list.append(os.path.join(war, item).replace('/','\\').replace('\exe','exe'))
                elif re.search(files_sql,item) != None and re.search(return_sql,item) == None and re.search(bak_sql,item) == None and re.search(select_sql,item) == None:
                    if os.path.exists(task_sql):
                        task_path = os.path.join(task_sql, item).replace('\\','/')
                    else:
                        os.makedirs(task_sql)
                        task_path = os.path.join(task_sql, item).replace('\\','/')
                    exe_list.append(os.path.join(sql, item).replace('/','\\').replace('\exe','exe'))
                elif re.search(return_sql,item)!=None:
                    if os.path.exists(task_back):
                        task_path = os.path.join(task_back, item).replace('\\','/')
                    else:
                        os.makedirs(task_back)
                        task_path = os.path.join(task_back, item).replace('\\','/')
                    exe_list.append(os.path.join(sql_back, item).replace('/','\\').replace('\exe','exe'))
                elif re.search(bak_sql,item)!=None:
                    if os.path.exists(task_bak):
                        task_path = os.path.join(task_bak, item).replace('\\','/')
                    else:
                        os.makedirs(task_bak)
                        task_path = os.path.join(task_bak, item).replace('\\','/')
                    exe_list.append(os.path.join(sql_bak, item).replace('/','\\').replace('\exe','exe'))
                elif re.search(select_sql,item)!=None:
                    if os.path.exists(task_select):
                        task_path = os.path.join(task_select, item).replace('\\','/')
                    else:
                        os.makedirs(task_select)
                        task_path = os.path.join(task_select, item).replace('\\','/')
                    exe_list.append(os.path.join(sql_select, item).replace('/','\\').replace('\exe','exe'))
                elif re.search(file_phone,item)!=None:
                    if os.path.exists(task_phone):
                        task_path = os.path.join(task_phone, item).replace('\\','/')
                    else:
                        os.makedirs(task_phone)
                        task_path = os.path.join(task_phone, item).replace('\\','/')
                    exe_list.append(os.path.join(phone, item).replace('/','\\').replace('\exe','exe'))

                #### 复制程序
                if re.search('_SP',task) !=None:
                    if re.search(files_auto_ins,item) == None and re.search(files_auto_exe,item) == None and re.search(files_checklist,item) == None and re.search('内部测试',item) == None:
                        shutil.copyfile(home_path,task_path)
                        i = i + 1
                        print(str(i) + '）' + '已收纳：' + item)
                    else:
                        print(item + ' 为非必要文件，已跳过！！！')
                else:
                    if re.search(files_auto_ins,item) == None and re.search(files_auto_exe,item) == None and re.search(files_checklist,item) == None:
                        shutil.copyfile(home_path,task_path)
                        i = i + 1
                        print(str(i) + '）' + '已收纳：' + item)
                    else:
                        print(item + ' 为非必要文件，已跳过！！！')                    
    return exe_list,src_list

############ 构造写docx函数 ###############
dt = datetime.datetime.now().strftime("%Y%m%d")[3:]

def write_prc_docx(file_prc,file_bf,document):
    p = document.add_paragraph('', style='List Number')
    r0 = p.add_run('将 ')
    r1 = p.add_run(file_prc)
    r2 = p.add_run(' 备份为 ')
    r3 = p.add_run(file_bf)
    r4 = p.add_run('''
注:第一步：在PLSQL里打开一个SQL Windows窗口，并复制粘贴备份的存储过程名；
第二步：选中该存储过程名字，然后点击鼠标右键；
第三步：在打开的菜单选项中找到【Edit】或者【编辑】并点击；
第四步：在新打开的窗口中，找到选中的存储过程名，并将其替换为要备份的存储过程名，然后执行【F8】''')
    B_list = [r1,r3]
    for B in B_list:
        B.italic = False
        B.font.size = Pt(11)
        B.font.highlight_color=WD_COLOR_INDEX.YELLOW
    R_list = [r0,r2,r4]
    for R in R_list:
        R.italic = False
        R.font.size = Pt(11)
        R.font.color.rgb = RGBColor(255,0,0)
    p.paragraph_format.left_indent = Pt(20)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST

def write_prc_rt_docx(file_bf,file_prc,document1):
    p = document1.add_paragraph('将 '+file_bf+' 还原为 '+file_prc, style='List Number')
    r0 = p.add_run('\r注:第一步：在PLSQL里打开一个SQL Windows窗口，写下存储过程名：')
    r1 = p.add_run(file_bf)
    r2 = p.add_run('''
第二步：选中该存储过程名字，然后点击鼠标右键；
第三步：在打开的菜单选项中找到【Edit】或者【编辑】并点击；
第四步：在新打开的窗口中，找到选中的存储过程名，并将''')
    r3 = p.add_run(file_bf)
    r4 = p.add_run('改成为')
    r5 = p.add_run(file_prc)
    r6 = p.add_run('后，\r然后执行【F8】；')
    B_list = [r1,r3,r5]
    for B in B_list:
        B.italic = False
        B.font.size = Pt(11)
        B.font.highlight_color=WD_COLOR_INDEX.YELLOW#黄色背景
    R_list = [r0,r2,r4,r6]
    for R in R_list:
        R.italic = False
        R.font.size = Pt(11)
        R.font.color.rgb = RGBColor(255,0,0)#红色字体
    p.paragraph_format.left_indent = Pt(20)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST

def write_dl_BF_docx(file_bf,document2):
    p = document2.add_paragraph('drop procedure '+file_bf, style=None)
    p.paragraph_format.left_indent = Pt(20)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST


def write_check_docx(check_list,document,p):
    for check_text in check_list:
        r = p.add_run(check_text+'\r')
        r.italic = False
        r.font.size = Pt(11)
        r.font.highlight_color=WD_COLOR_INDEX.YELLOW
    p.paragraph_format.left_indent = Pt(20)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST

################ 构造读取预置表函数 #############
def read_book(excel_path,sheet):
    workbook1=load_workbook(excel_path)
    sheet=workbook1[sheet]
    max_row=sheet.max_row
    for row in range(2,max_row+1):
        sub_data={}
        sub_data[title1]=sheet.cell(row,1).value
        sub_data[title2]=sheet.cell(row,2).value
        sub_data[title3]=sheet.cell(row,3).value
        if int(sub_data[title3]) >= 0:
            model_list.append(sub_data)

def read_ebook(excel_check_path,sheet):
    workbook1=load_workbook(excel_check_path)
    sheet=workbook1[sheet]
    max_row=sheet.max_row
    for row in range(2,max_row+1):
        sub_data={}
        sub_data[title4]=sheet.cell(row,1).value
        sub_data[title5]=sheet.cell(row,2).value
        sub_data[title6]=sheet.cell(row,3).value
        if sub_data[title5] == '海鹰':
            HY_check_list.append(sub_data)
        else:
            HSH_check_list.append(sub_data)
    return HY_check_list,HSH_check_list

def check_combine(check_result,check_list,combine_list):
    if check_result != [] and check_list != []:
        for w_pro_name in check_result:
            for check in check_list:
                pro_name = ''.join(check.get(title4))
                check_run = ''.join(check.get(title6))
                if pro_name == w_pro_name:
                    combine_list.append(check_run)
                    break
    return combine_list

################ 构造合并docx文档函数 ###########
def combine_all_docx(filename_master,files_list):
    number_of_sections=len(files_list)
    master = Document_compose(filename_master)
    composer = Composer(master)
    print(files_list[0])
    for i in range(1, number_of_sections):
        doc_temp = Document_compose(files_list[i])
        print(files_list[i])
        composer.append(doc_temp)
    composer.save(task_reldoc+'/'+files_auto_ins+'.docx')#保存为指定文件名

################ 清理异常文件并调用扫描函数 ###################
if os.path.exists(home+file_source+'/PRC_BF.docx'): 
    os.remove(home+file_source+'/PRC_BF.docx')
if os.path.exists(home+file_source+'/PRC_RT.docx'):
    os.remove(home+file_source+'/PRC_RT.docx')
if os.path.exists(home+file_source+'/PRC_DL.docx'):
    os.remove(home+file_source+'/PRC_DL.docx')
if os.path.exists(home+file_source+'/CHECK.docx'):
    os.remove(home+file_source+'/CHECK.docx')

fileListFunc([home+file_source],exe_list,src_list)

################# 生成程序完成清单 ##################
if os.path.exists(task_indoc):
    for name in os.listdir(package_model):
            auto_exe_model = os.path.join(package_model, name).replace('\\','/')
            auto_exe = os.path.join(task_indoc, name).replace('\\','/')
            if len(name.split(".")) == 2:
                    if re.search(files_auto_exe,name.split(".")[0]) != None:
                            wb = Workbook()
                            wb = load_workbook(auto_exe_model)
                            ws = wb["Sheet1"]
                            # 激活 worksheet
                            ws = wb.active
                            # 数据可以直接分配到单元格中
                            ws['B2'] = task
                            ws['E2'] = tester
                            ws['G2'] = task_type
                            # 可以附加行，从第一列开始附加
                            write_list = src_list+exe_list
                            j = 0
                            for write in write_list:
                                    j = j+1
                                    ws.append([j,write,'源代码' if re.search('src',write) !=None else '执行码',devlpor,teamer,tester])
                            # 保存文件
                            wb.save(auto_exe)
                            print('程序完成清单已生成！！！')
else:
    print('indoc目录不存在，程序完成清单生成已跳过，请检查是否遗漏项目文档后重新组包！！！')

############ 判断组包类型并根据不同包生成安装手册 #################
############ 构造写文件函数 #################
def check_and_change(document, replace_dict, new_file):#docx分为段落里的run和表格里的cell两部分逐个替换
    j=0
    k=0#敏感词计数
    ###check敏感词
    for para in document.paragraphs:
        for i in range(len(para.runs)):
            for key, value in replace_dict.items():
                j = j + para.runs[i].text.count(key)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in replace_dict.items():
                    k = k + cell.text.count(key)
    if j+k>0:#若j+k大于零说明有敏感词
        ###change敏感词
        for para in document.paragraphs:
            for run in para.runs:
                for key, value in replace_dict.items():
                    if key in run.text:
                        run.text = run.text.replace(key, value)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in replace_dict.items():
                        if key in cell.text:
                            cell.text = cell.text.replace(key, value)
        document.save(new_file)#保存新文件
    return document

######### 根据不同版本选择不同模式生成安装手册###
# 判断为客户端版本
if task_type == 'MLIFE-PHONE':
    auto_ipa_list = []
    for name in os.listdir(task_phone):
        if re.search('.ipa',name) != None:
            auto_ipa_list.append(name)
    if len(auto_ipa_list) == 1:
        auto_ipa = ''.join(auto_ipa_list)
    else:
        auto_ipa = 'Mlife.ipa'
    replace_dict = {'【编辑日期】':datetime.datetime.now().strftime("%Y{y}%m{m}%d{d}").format(y='年', m='月', d='日'),'【版本号】':task,'【新建日期】':datetime.datetime.now().strftime("%Y-%m-%d"),'【测试】':tester,'【安装包名】':auto_ipa,}
    for name in os.listdir(package_model):
        auto_install_model = os.path.join(package_model, name).replace('\\','/')
        auto_install = os.path.join(task_reldoc, name).replace('\\','/')
        if len(name.split(".")) == 2:
            if re.search(files_auto_ins,name.split(".")[0]) != None:
                document2 = docx.Document(auto_install_model)
                document2 = check_and_change(document2, replace_dict, auto_install)
                print('安装手册已生成。')

#判断为MLIFE-APPS版本
elif task_type == 'MLIFE-APPS':
    #读取目录获取执行码清单备用
    sql_dir = code_dir+code_list[0]
    war_dir = code_dir+code_list[1]
    rtsql_dir = code_dir+code_list[2]
    find_file(slist,sql_dir,['backsql','exesql'])
    find_file(rlist,rtsql_dir,['returnsql','returnsql_check'])
    find_file(wlist,war_dir,[])
    if slist !=[]:
        for prc in slist:
            if re.search(r'PRC|prc',prc):
                plist.append(prc)
    if plist !=[]:
        document = Document_compose()
        document1 = Document_compose()
        document2 = Document_compose()
        p0 = document.add_heading('备份存储过程：', 5)
        p1 = document1.add_heading('回滚存储过程：', 5)
        p2 = document2.add_heading('删除备份存储过程：', 5)
        p0.paragraph_format.left_indent = Pt(10)
        p1.paragraph_format.left_indent = Pt(10)
        p2.paragraph_format.left_indent = Pt(10)
        for file in plist:
            file_prc = file.split('.')[0]
            if len(file_prc)>22:
                file_bf = file.split('.')[0][:22]+str(plist.index(file))+'_'+str(dt)+'B'
            else:
                file_bf = file.split('.')[0]+'_'+str(dt)+'BF'
            write_prc_docx(file_prc,file_bf,document)
            write_prc_rt_docx(file_bf,file_prc,document1)
            write_dl_BF_docx(file_bf,document2)
            
        #保存.docx文档
        document.save(home+file_source+'/PRC_BF.docx')
        document1.save(home+file_source+'/PRC_RT.docx')
        document2.save(home+file_source+'/PRC_DL.docx')
        wlist.append('PRC_BF')
        wlist.append('PRC_RT')
        wlist.append('PRC_DL')

    #读取预置表，获取有序的模版列表
    read_book(excel_path,task_type.split("-")[0])#读取指定sheet页
    for mod in model_list:
        mod_name = ''.join(mod.get(title2))
        use_num = ''.join(str(mod.get(title3)))
        if mod_name in wlist:
            pro_t_name = mod_name.split('.')[0]
            model_check_list.update({pro_t_name:use_num})
    n=zip(model_check_list.values(),model_check_list.keys())
    n_e = sorted(n,reverse=True)
    for e in n_e:
        check_result.append(e[1])

    ####  读取excel表生成验证清单   ######
    read_ebook(excel_check_path,task_type.split("-")[0])#读取指定sheet页
    check_combine(check_result,HY_check_list,HY_combine_list)
    check_combine(check_result,HSH_check_list,HSH_combine_list)
    document3 = Document_compose()
    p3 = document3.add_heading('1）正向验证：', 5)
    p4 = document3.add_paragraph('''部署完成后，打开浏览器，输入如下地址查看服务是否正常访问并反显访问GCSC访问地址：\r''', style=None)
    if HSH_combine_list !=[]:
        p9 = document3.add_paragraph('黑山扈：\r', style=None)
        write_check_docx(HSH_combine_list,document3,p9)
    if HY_combine_list !=[]:
        p5 = document3.add_paragraph('海鹰：\r', style=None)
        write_check_docx(HY_combine_list,document3,p5)
    p6 = document3.add_heading('2）回退验证：', 5)
    p7 = document3.add_paragraph('''部署完成后，打开浏览器，输入如下地址查看服务是否正常访问并反显访问GCSC访问地址：\r''', style=None)
    if HSH_combine_list !=[]:
        p10 = document3.add_paragraph('黑山扈：\r', style=None)
        write_check_docx(HSH_combine_list,document3,p10)
    if HY_combine_list !=[]:
        p8 = document3.add_paragraph('海鹰：\r', style=None)
        write_check_docx(HY_combine_list,document3,p8)
    
    document3.save(home+file_source+'/CHECK.docx')
    wlist.append('check_start')
    if HSH_combine_list !=[] or HY_combine_list !=[]:
        wlist.append('CHECK')
    wlist.append('end.war')
    #读取预置表，获取有序的模版列表
    read_book(excel_path,task_type.split("-")[0])#读取指定sheet页
    for mod in model_list:
        file_path = ''.join(mod.get(title1))#指定表头名称，下同
        mod_name = ''.join(mod.get(title2))
        use_num = ''.join(str(mod.get(title3)))
        if mod_name in wlist:
            file = home + file_path + '/' + mod_name + '.docx'#限制文档格式为docx
            model_use_list.update({file:use_num})
    m=zip(model_use_list.values(),model_use_list.keys())
    m_d = sorted(m,reverse=True)
    for d in m_d:
        result.append(d[1])
    #调用合并文档函数合并生成安装手册
    combine_all_docx(result[0],result)
    #动态替换生成的安装手册中所需修改的参数
    replace_dict = {'【编辑日期】':datetime.datetime.now().strftime("%Y{y}%m{m}%d{d}").format(y='年', m='月', d='日'),'【版本号】':task,'【新建日期】':datetime.datetime.now().strftime("%Y-%m-%d"),'【测试】':tester,'【执行sql】':'\n'.join(slist),'【回滚sql】':'\n'.join(rlist)}
    for name in os.listdir(task_reldoc):
        auto_install = os.path.join(task_reldoc, name).replace('\\','/')
        if len(name.split(".")) == 2:
            if re.search(files_auto_ins+'.docx',name) != None:
                document = docx.Document(auto_install)
                document = check_and_change(document, replace_dict, auto_install)
                print('安装手册已生成。')
    if os.path.exists(home+file_source+'/PRC_BF.docx'): 
        os.remove(home+file_source+'/PRC_BF.docx')
    if os.path.exists(home+file_source+'/PRC_RT.docx'):
        os.remove(home+file_source+'/PRC_RT.docx')
    if os.path.exists(home+file_source+'/PRC_DL.docx'):
        os.remove(home+file_source+'/PRC_DL.docx')
    if os.path.exists(home+file_source+'/CHECK.docx'):
        os.remove(home+file_source+'/CHECK.docx')
    if task_type == "MLIFE-APPS":
        # 复制checklist并动态修改系统参数，不修改具体检查项
        for name in os.listdir(package_model):
            check_model = os.path.join(package_model, name).replace('\\','/')
            auto_checklist = os.path.join(task_indoc, name).replace('\\','/')
            if len(name.split(".")) == 2:
                if re.search(files_checklist,name.split(".")[0]) != None:
                    wb = Workbook()
                    wb = load_workbook(check_model)
                    ws = wb["Sheet1"]
                    ws = wb.active
                    ws['B2'] = task
                    ws['E2'] = devlpor
                    ws['G2'] = datetime.datetime.now().strftime("%Y/%m/%d")
                    ws['E34'] = teamer
                    ws['G34'] = datetime.datetime.now().strftime("%Y/%m/%d")
                    wb.save(auto_checklist)
                    print('缤纷生活checklist已生成，请参照检查项逐一检查！！！')

#调用关闭窗口倒计时
Windows_close(300)
>>>>>>> e03e7e280d6a4bc9498875f2b37d0f49ef931195
