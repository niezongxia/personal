# coding=utf-8

import os
import csv
import docx
import time
import codecs

#路径
home = "./"# 放了一些docx 文件
home = "./"# 生成新文件后的存放地址
home = './'#字典路径

#字典初始化
replace_dict = {}

def csv_dict(home):#csv转换成字典
    for fn in os.listdir(home):
        dict_file = home+fn
        if len(fn.split(".")) == 2:
            if fn.split(".")[1] == 'csv':
                with codecs.open(dict_file,'r',encoding='gbk') as f:
                    reader = csv.DictReader(f)
                    for row in reader:
                        csv_dict = dict(row)
                    replace_dict.update(csv_dict)

def check_and_change(document, replace_dict, new_file):#docx分为段落里的run和表格里的cell两部分逐个替换
    j=0
    k=0#敏感词计数
    ###check敏感词
    for para in document.paragraphs:
        for i in range(len(para.runs)):
            for key, value in replace_dict.items():
                j=j+para.runs[i].text.count(key)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in replace_dict.items():
                    k=k+cell.text.count(key)
    if j+k>0:#若j+k大于零说明有敏感词
        ###change敏感词
        for para in document.paragraphs:
            for run in para.runs:
                for key, value in replace_dict.items():
                    if key in run.text:
                        run.text = run.text.replace(key, value)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in replace_dict.items():
                        if key in cell.text:
                            cell.text = cell.text.replace(key, value)
        document.save(new_file)#保存新文件
        print("测试报告： "+new_file)
    return document

def Windows_close(close_times):
    for x in range(close_times,-1,-1):
        mystr = "程序执行完毕！窗口将在" + str(x) + "秒后关闭！！！"
        print(mystr,end = "")
        print("\b" * (len(mystr)*2),end = "",flush=True)
        time.sleep(1)

def main():
    csv_dict(home)
    for name in os.listdir(home):
        old_file = home + name
        new_name = name.replace('【模板】','')
        new_file = home + new_name
        if len(name.split(".")) == 2:
            if name.split(".")[1] == 'docx':
                print("模版文件： "+old_file)
                document = docx.Document(old_file)
                document = check_and_change(document, replace_dict, new_file)
    print('程序执行完毕。')
    Windows_close(300)
if __name__ == '__main__':
    main()
